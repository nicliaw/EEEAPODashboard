import pandas as pd
import os
import matplotlib
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    send_from_directory,
)
from flask_sqlalchemy import SQLAlchemy
from docx import Document
from docx.shared import Inches
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

matplotlib.pyplot.switch_backend("Agg")

basedir = os.path.abspath(os.path.dirname(__file__))

app = Flask(__name__)
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///" + os.path.join(
    basedir, "database/comments.db"
)
db = SQLAlchemy(app)
comments_data = {}


# Comment Model
class Comment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    file = db.Column(db.Text)
    name = db.Column(db.String(80))
    email = db.Column(db.String(120))
    comment_text = db.Column(db.Text)
    timestamp = db.Column(
        db.String(19), default=lambda: datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    )


@app.route("/", methods=["GET", "POST"])
def index():
    academic_years = get_academic_years()
    first_year = academic_years[0]
    excel_files = get_excel_files_by_ay(academic_years, first_year)
    
    if request.method == "POST":
        return dashboard()

    if excel_files:
        first_file = excel_files[0]
        return dashboard(first_file)

    return render_template("dashboard.html", excel_files=excel_files)


@app.route("/dashboard", methods=["POST"])
def dashboard(selected_file=None):
    if not selected_file:
        selected_file = request.form.get("file")
    selected_cohort = request.form.get("cohort", "All")
    academic_years = get_academic_years()
    selected_academic_year = request.form.get("ay")
    if selected_academic_year is None and academic_years:
        selected_academic_year = academic_years[0]

    excel_files = get_excel_files_by_ay(academic_years, selected_academic_year)
    # Parse the course sizes from total.txt
    course_sizes = parse_total_file(selected_academic_year)


    file_path = os.path.join("excel", selected_file) if selected_file else None

    # Read the Excel file
    df = pd.read_excel(file_path, sheet_name="Form1")

    # Apply filter based on selected cohort
    if selected_cohort != "All":
        df = df[df["Cohort"] == selected_cohort]

    # Deduplicate column names
    df.columns = list(deduplicate_columns(df.columns))

    # Pass the entire original table to the template
    original_column_names = df.columns.tolist()
    original_data = df.to_dict("records")

    # Extract columns for the "Extracted Table": the three columns after "Question"
    start_col_extracted = df.columns.get_loc("Question") + 1
    end_col_extracted = start_col_extracted + 3
    extracted_column_headers = df.columns[
        start_col_extracted:end_col_extracted
    ].tolist()

    # Extract the corresponding data for the extracted column headers for "Extracted Table"
    extracted_data = []
    for _, row in df.iterrows():
        extracted_row = {}
        for header in extracted_column_headers:
            extracted_row[header] = row[header]
        extracted_data.append(extracted_row)

    # Extract columns for the "Reformatted Table": from the column after "Cohort" to the "Question" column
    start_col_reformatted = df.columns.get_loc("Cohort") + 1
    end_col_reformatted = df.columns.get_loc("Question")
    reformatted_column_headers = df.columns[
        start_col_reformatted:end_col_reformatted
    ].tolist()

    # Prepare the data for "Reformatted Table"
    reformatted_data_input = []
    for _, row in df.iterrows():
        reformatted_row = {}
        for header in reformatted_column_headers:
            reformatted_row[header] = row[header]
        reformatted_data_input.append(reformatted_row)

    # Compute the reformatted data
    reformatted_data = compute_percentages_and_counts(
        reformatted_data_input, reformatted_column_headers
    )
    column_names = ["Statement/Question"] + sorted(reformatted_data.columns.tolist())
    data = [
        {"Statement/Question": column, **row.to_dict()}
        for column, row in reformatted_data.iterrows()
    ]

    # Generate the word cloud
    wordcloud_image_path = generate_word_cloud(file_path, extracted_column_headers[0])

    # Get comments for the selected file
    comments = Comment.query.filter(Comment.file == selected_file).all()

    course_code = extract_course_code(selected_file)
    course_size = course_sizes[course_code]
    num_respondents = len(df)  # Number of rows in the Excel file
    response_rate = (num_respondents / course_size) * 100

    return render_template(
        "dashboard.html",
        excel_files=excel_files,
        data=data,
        column_names=column_names,
        original_data=original_data,
        original_column_names=original_column_names,
        extracted_data=extracted_data,
        extracted_column_headers=extracted_column_headers,
        wordcloud_image_url=wordcloud_image_path,
        comments=comments,
        selected_file=selected_file,
        academic_years=academic_years,
        selected_cohort=selected_cohort,
        selected_academic_year=selected_academic_year,
        response_rate=round(response_rate, 2)
    )


@app.route("/comment", methods=["POST"])
def comment():
    if request.method == "POST":
        file = request.form.get("file")
        name = request.form.get("name")
        email = request.form.get("email")
        comment_text = request.form.get("comment")
        c = Comment(file=file, name=name, email=email, comment_text=comment_text)
        db.session.add(c)
        db.session.commit()
    comments = Comment.query.all()

    # Create a Pandas DataFrame to store the data
    data = {'Name': [], 'Email': [], 'Comment': []}

    for comment in comments:
        data['Name'].append(comment.name)
        data['Email'].append(comment.email)
        data['Comment'].append(comment.comment_text)

    df = pd.DataFrame(data)

    # Create a new Excel workbook and write the data
    output_filename = 'comments.xlsx'
    writer = pd.ExcelWriter(output_filename, engine='xlsxwriter')
    df.to_excel(writer, sheet_name='Sheet1', index=False)
    writer.close()  # Use save() to save and close the Excel file
    return redirect(url_for("index"))


def deduplicate_columns(columns):
    """Ensure all columns are unique, appending .1, .2, etc. to duplicates."""
    seen = {}
    for i, col in enumerate(columns):
        if col not in seen:
            seen[col] = 1
            yield col
        else:
            seen[col] += 1
            yield f"{col}.{seen[col] - 1}"


def get_excel_files():
    folder_path = os.path.join(os.path.dirname(__file__), "excel")
    return [file for file in os.listdir(folder_path) if file.endswith(".xlsx")]

def get_excel_files_by_ay(academic_years, selected_academic_year):
    excel_files = []
    for academic_year in academic_years:
        folder_path = os.path.join(os.path.dirname(__file__), f"excel/{academic_year}")
        if os.path.exists(folder_path):
            files = [os.path.join(academic_year, file) for file in os.listdir(folder_path) if file.endswith(".xlsx")]
            excel_files.extend(files)

    # Filter files by the selected academic year
    filtered_files = [file for file in excel_files if file.startswith(selected_academic_year)]
    return filtered_files


def get_academic_years():
    # Get a list of academic years (folder names) within the "excel" directory
    academic_years = []
    excel_dir = os.path.join(os.path.dirname(__file__), "excel")
    if os.path.exists(excel_dir):
        academic_years = [folder for folder in os.listdir(excel_dir) if os.path.isdir(os.path.join(excel_dir, folder))]
    return academic_years


def extract_course_code(selected_file):
    # The course code is the part between the last '/' and '.xlsx'
    course_code = selected_file.split('/')[-1].replace('.xlsx', '')
    return course_code
    
def parse_total_file(academic_year):
    total_file_path = os.path.join("excel", academic_year, "total.txt")
    course_sizes = {}
    
    with open(total_file_path, "r") as file:
        for line in file:
            parts = line.strip().split(":")
            if len(parts) == 2:
                course_code, size = parts[0].strip(), int(parts[1].strip())
                course_sizes[course_code] = size

    return course_sizes

def compute_percentages_and_counts(extracted_data, extracted_column_headers):
    reformatted_data = pd.DataFrame(index=extracted_column_headers)

    for response in [
        "Strongly Agree",
        "Agree",
        "Neutral",
        "Disagree",
        "Strongly Disagree",
    ]:
        response_counts = [
            pd.Series([row[header] for row in extracted_data])
            .value_counts()
            .get(response, 0)
            for header in extracted_column_headers
        ]
        total_responses = len(extracted_data)

        # Compute the percentage, ensuring we avoid a division by zero
        reformatted_data[response] = [
            round((count / total_responses) * 100, 2) if total_responses > 0 else 0
            for count in response_counts
        ]
        reformatted_data[response + " Count"] = response_counts

    return reformatted_data


def generate_word_cloud(file_path, selected_column):
    if file_path and os.path.isfile(file_path):
        # Read the Excel file
        df = pd.read_excel(file_path, sheet_name="Form1")

        # Extract the text data from the selected column
        extracted_text = " ".join(df[selected_column].astype(str))

        # Create a WordCloud object
        wordcloud = WordCloud(
            width=800, height=400, background_color="#f8f9fa"
        ).generate(extracted_text)

        # Generate the word cloud image
        plt.figure(figsize=(10, 5))
        plt.imshow(wordcloud, interpolation="bilinear")
        plt.axis("off")

        # Save the word cloud image to a file in the static folder
        wordcloud_image_path = "static/wordcloud.png"
        plt.savefig(wordcloud_image_path, facecolor="#f8f9fa")

        # Return the file path of the word cloud image
        return wordcloud_image_path
    else:
        return None

@app.route("/generate_report", methods=["POST"])
def generate_report():
    selected_file = request.form.get("selected_file")
    if not selected_file:
        return "No file selected", 400

    file_path = os.path.join("excel", selected_file)
    if not os.path.exists(file_path):
        return "File not found", 404

    # Read the "Form1" sheet from the selected Excel file
    df = pd.read_excel(file_path, sheet_name="Form1")

    # Deduplicate column names
    df.columns = list(deduplicate_columns(df.columns))

    # Extract columns for the "Reformatted Table": from the column after "Cohort" to the "Question" column
    start_col_reformatted = df.columns.get_loc("Cohort") + 1
    end_col_reformatted = df.columns.get_loc("Question")
    reformatted_column_headers = df.columns[
        start_col_reformatted:end_col_reformatted
    ].tolist()

    # Prepare the data for "Reformatted Table"
    reformatted_data_input = []
    for _, row in df.iterrows():
        reformatted_row = {}
        for header in reformatted_column_headers:
            reformatted_row[header] = row[header]
        reformatted_data_input.append(reformatted_row)

    # Compute the reformatted data
    reformatted_data = compute_percentages_and_counts(
        reformatted_data_input, reformatted_column_headers
    )

    # Generate the Word document
    doc = Document()
    doc.add_heading("Reformatted Table", 0)

    # Add the reformatted table to the document
    table = doc.add_table(
        rows=reformatted_data.shape[0] + 2, cols=reformatted_data.shape[1] + 1
    )

    # Set headers for the table
    table.cell(0, 0).text = "Question"
    for i, column in enumerate(reformatted_data.columns, start=1):
        table.cell(0, i).text = column

    # Fill in the table data
    for i, (index, row) in enumerate(reformatted_data.iterrows(), start=1):
        table.cell(i, 0).text = index
        for j, header in enumerate(reformatted_data.columns, start=1):
            table.cell(i, j).text = str(row[header])

    # Save the document
    doc.save("static/report.docx")
    return send_from_directory("static", "report.docx", as_attachment=True)

# Function to send an email
def send_email(subject, body, to_email):
    # Your Gmail account credentials
    email = 'your_email@gmail.com'
    password = 'your_password'

    # Create the email message
    msg = MIMEMultipart()
    msg['From'] = email
    msg['To'] = to_email
    msg['Subject'] = subject

    # Attach the email body
    msg.attach(MIMEText(body, 'plain'))

    try:
        server = smtplib.SMTP('smtp-mail.outlook.com', 587)
        server.starttls()
        server.login(email, password)

        # Send the email
        text = msg.as_string()
        server.sendmail(email, to_email, text)

        # Close the server connection
        server.quit()
        return True
    except Exception as e:
        return str(e)


if __name__ == "__main__":
    app.run(debug=True)
